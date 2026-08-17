"""Build the living project progress report as a polished Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Competitive_Connectomes_Progress_Report.docx"
FIGURES = ROOT / "figures"

BLUE = "2E5D7B"
DARK_BLUE = "173A52"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "68717A"
WHITE = "FFFFFF"
BLACK = "1A1A1A"
GOLD = "A77A2D"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)

    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")


def style_run(run, size=11, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text: str, *, bold_lead: str | None = None, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.keep_together = True
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        style_run(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        style_run(rest, italic=italic)
    else:
        run = paragraph.add_run(text)
        style_run(run, italic=italic)
    return paragraph


def add_bullet(doc, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.167
    style_run(paragraph.add_run(text))
    return paragraph


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_properties.append(cant_split)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    style_run(paragraph.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
    style_run(paragraph.add_run(text), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_result_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(paragraph.add_run(header), bold=True, color=DARK_BLUE, size=10)
    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row_values, strict=True)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_run(paragraph.add_run(value), size=10)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc, filename: str, caption: str, width=6.2):
    path = FIGURES / filename
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(3)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    style_run(caption_paragraph.add_run(caption), size=9, italic=True, color=MID_GRAY)


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    style_run(run, size=9, color=MID_GRAY)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167


def build_document() -> None:
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(header.add_run("COMPETITIVE CONNECTOMES AND ADAPTIVE FLEXIBILITY"), size=8.5, bold=True, color=MID_GRAY)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(26)
    title.paragraph_format.space_after = Pt(5)
    style_run(title.add_run("Research Progress Report"), size=25, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    style_run(
        subtitle.add_run("Competitive Connectomes and Adaptive Flexibility"),
        size=15,
        color=BLUE,
    )

    metadata = [
        ("Project stage", "EEG estimator comparison complete; predictive drift correction next"),
        ("Report status", "Living document — update after each milestone"),
        ("Data scope", "Released Luppi et al. data plus OpenNeuro ds004295 EEG"),
        ("Last updated", date(2026, 8, 11).strftime("%d %B %Y")),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        style_run(paragraph.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
        style_run(paragraph.add_run(value))

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    ppr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    ppr.append(borders)

    add_callout(
        doc,
        "Current evidence",
        "For the released single subject, cooperative and competitive effective interactions jointly contribute to reproducing empirical static functional connectivity. Their signs, strengths and anatomical assignments form an interdependent fitted organization.",
    )

    add_heading(doc, "1. Executive summary", 1)
    add_body(
        doc,
        "This project began as a reproducibility study of the signed Hopf whole-brain model released with Luppi et al. It has now completed the core single-subject replication and a set of controlled perturbation experiments designed to determine why the fitted signed network succeeds. The work was completed using public data and an independent Python/C++ workflow without MATLAB."
    )
    add_body(
        doc,
        "The signed model consistently reproduced empirical functional connectivity better than the cooperative-only model. Subsequent null experiments showed that this advantage was not explained solely by the presence of negative weights, the overall weight distribution or the anatomical connectivity mask. Cooperative strengths, competitive strengths and their anatomical sign assignment each contributed to the fitted network’s performance."
    )
    add_callout(
        doc,
        "Interpretive boundary",
        "These results concern a frozen model fitted to one released subject. Simulation seeds quantify stochastic robustness, not variation across people. The experiments support dependence within this fitted configuration, not uniqueness among all possible reoptimized solutions.",
        fill=LIGHT_GRAY,
    )

    doc.add_page_break()
    add_heading(doc, "2. Scientific framing", 1)
    add_heading(doc, "2.1 Research question", 2)
    add_body(
        doc,
        "How do biologically inferred cooperative and competitive effective interactions jointly organize brain dynamics, and how might their balance support stable exploitation and adaptive exploration?"
    )
    add_heading(doc, "2.2 Working hypothesis", 2)
    add_body(
        doc,
        "Empirical brain dynamics emerge from the complementary spatial organization of cooperative and competitive effective interactions. Disrupting either component should impair model fit, while their relative balance may regulate the trade-off between stable maintenance and adaptive switching."
    )
    add_heading(doc, "2.3 Terminology", 2)
    add_result_table(
        doc,
        ["Term", "Meaning in this project"],
        [
            ["Structural connectivity", "The anatomical connection mask and strengths supplied to the model."],
            ["Functional connectivity", "Correlations between regional BOLD time series."],
            ["Effective connectivity", "Model-inferred directional influences used to reproduce empirical dynamics."],
            ["Cooperative interaction", "A positive fitted effective weight."],
            ["Competitive interaction", "A negative fitted effective weight; not necessarily a direct long-range inhibitory synapse."],
        ],
        [2300, 7060],
    )

    doc.add_page_break()
    add_heading(doc, "3. Phase 1 — Reproducible computational foundation", 1)
    add_body(doc, "Status: complete.", bold_lead="Status:")
    for item in [
        "Created and verified an isolated Python 3.12.9 environment.",
        "Cloned the authors’ repository and recorded commit da592aab6784db5c6c59f29f6bcb2b3743f1afd7.",
        "Built and smoke-tested the C++ Hopf simulator and optimizer without MATLAB.",
        "Loaded the released 100-region structural-connectivity matrix and 1,189-timepoint BOLD data.",
        "Translated regional-frequency extraction from MATLAB to Python and matched the released reference exactly.",
        "Added optional random-seed control while preserving the upstream default of seed 42.",
        "Established automated tests for loading, frequency extraction, null-model preservation and stochastic simulation.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Reproducibility note",
        "All upstream modifications are documented separately. The original default behaviour remains available, and the seed extension is used only for repeated stochastic validation.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "4. Phase 2 — Core replication", 1)
    add_heading(doc, "4.1 Data and fitted models", 2)
    add_body(
        doc,
        "The released single-subject BOLD data were used to derive 100 characteristic regional oscillation frequencies. Two Hopf models were optimized with otherwise matched settings: a cooperative-only model constrained to non-negative generative weights and a signed model permitted to infer both positive and negative effective weights."
    )
    add_result_table(
        doc,
        ["Property", "Cooperative-only", "Signed"],
        [
            ["Positive directed weights", "1,464", "1,493"],
            ["Negative directed weights", "0", "767"],
            ["Total non-zero directed weights", "1,464", "2,260"],
            ["Negative fraction of non-zero weights", "0%", "33.94%"],
        ],
        [4200, 2580, 2580],
    )

    add_heading(doc, "4.2 Frozen-model stochastic validation", 2)
    add_body(
        doc,
        "The optimized matrices were frozen and forward-simulated across 30 matched random-noise seeds. This separated stochastic simulation variability from optimization and prevented the comparison from depending on a single random realization."
    )
    add_result_table(
        doc,
        ["Frozen model", "Mean FC correlation", "SD", "Mean FC MAE"],
        [
            ["Cooperative-only", "0.482", "0.030", "0.161"],
            ["Signed", "0.677", "0.025", "0.130"],
        ],
        [3600, 2100, 1560, 2100],
    )
    add_body(
        doc,
        "The signed model performed better in all 30 matched simulations. Its mean FC-correlation advantage was 0.196, indicating that the improvement was robust to the model’s stochastic neural-noise process."
    )
    add_figure(
        doc,
        "frozen_gc_stochastic_evaluation.png",
        "Figure 1. Frozen cooperative-only and signed models across 30 matched stochastic seeds.",
    )

    add_heading(doc, "5. Expanded Phase 2 — Mechanistic null experiments", 1)
    add_heading(doc, "5.1 Whole-network weight-placement null", 2)
    add_body(
        doc,
        "Reciprocal fitted weight pairs were permuted among anatomically occupied edges. The exact weight multiset, sign counts, density, anatomical mask and reciprocal-pair asymmetry were preserved, while the anatomical placement of the complete fitted weight set was disrupted."
    )
    add_result_table(
        doc,
        ["Configuration", "Mean FC correlation"],
        [
            ["Original fitted placement", "0.682"],
            ["Mean of 100 shuffled placements", "0.067"],
            ["Best shuffled placement", "0.116"],
        ],
        [6500, 2860],
    )
    add_body(
        doc,
        "No shuffled network matched the original (one-sided permutation p = 0.0099). The complete fitted weight organization was therefore highly location-dependent. Because positive and negative weights moved together, this null did not isolate either sign individually."
    )
    add_figure(
        doc,
        "signed_weight_placement_null.png",
        "Figure 2. Original fitted weight placement compared with 100 topology-constrained shuffled networks.",
    )

    add_heading(doc, "5.2 Sign-specific strength organization and ablation", 2)
    add_body(
        doc,
        "Negative strengths were shuffled only among the 767 already-negative locations. As a balanced control, exactly 767 positive weights were selected to approximately match the negative magnitude distribution and shuffled only among those positive locations. Every non-targeted weight remained fixed."
    )
    add_result_table(
        doc,
        ["Condition", "Mean FC correlation", "Loss from original"],
        [
            ["Original signed network", "0.682", "—"],
            ["Negative-strength shuffle", "0.562", "0.119"],
            ["Matched positive-strength shuffle", "0.604", "0.078"],
            ["Negative weights removed", "0.443", "0.239"],
            ["Negative weights flipped positive", "0.344", "0.337"],
        ],
        [4800, 2280, 2280],
    )
    add_body(
        doc,
        "Both strength shuffles significantly impaired fit (p = 0.0099 for each null). The result supports location-specific organization of both cooperative and competitive strengths. Removing negative interactions damaged performance, and changing them to positive was more damaging still, indicating that their competitive sign—not merely their magnitudes and locations—was important."
    )
    add_body(
        doc,
        "The negative-strength shuffle caused a larger average loss than the matched positive shuffle in this subject. This is not evidence that competition is universally more important: the analysis uses one subject and a magnitude-matched subset rather than all positive weights."
    )
    add_figure(
        doc,
        "sign_specific_organization.png",
        "Figure 3. Negative-interaction ablations and matched sign-specific strength perturbations.",
    )

    add_heading(doc, "5.3 Anatomical sign-map null", 2)
    add_body(
        doc,
        "Every directed connection retained its exact absolute magnitude and anatomical location. Reciprocal sign-pattern pairs were redistributed among occupied region pairs while preserving the total positive/negative counts and the numbers of cooperative–cooperative, mixed-sign and competitive–competitive reciprocal patterns."
    )
    add_result_table(
        doc,
        ["Sign map", "Mean FC correlation"],
        [
            ["Original fitted sign map", "0.682"],
            ["Mean of 100 randomized sign maps", "0.040"],
            ["Best randomized sign map", "0.151"],
        ],
        [6500, 2860],
    )
    add_body(
        doc,
        "None of the randomized maps approached the original (one-sided permutation p = 0.0099). Thus, topology, magnitudes and overall cooperative–competitive balance were insufficient by themselves: the fitted signs also had to be assigned to appropriate anatomical locations."
    )
    add_figure(
        doc,
        "sign_map_null.png",
        "Figure 4. Original anatomical sign map compared with 100 magnitude-preserving randomized sign maps.",
    )

    add_heading(doc, "6. Integrated interpretation", 1)
    add_body(
        doc,
        "The completed experiments support an evidence ladder rather than a single all-or-nothing claim. Each experiment holds different properties constant and isolates a different component of the fitted organization."
    )
    add_result_table(
        doc,
        ["Evidence step", "Supported interpretation"],
        [
            ["Signed > cooperative-only", "Allowing cooperative and competitive effective interactions improves static FC fit."],
            ["Frozen-seed validation", "The advantage is not dependent on one stochastic forward simulation."],
            ["Whole-weight shuffle", "The complete fitted weight set is anatomically organized."],
            ["Within-sign shuffles", "Both cooperative and competitive strength-to-location mappings carry information."],
            ["Removal and sign flip", "Competitive interactions must be present and retain their negative sign."],
            ["Sign-map null", "The anatomical allocation of cooperative and competitive roles contributes beyond fixed topology and magnitudes."],
        ],
        [3050, 6310],
    )
    add_callout(
        doc,
        "Bounded conclusion",
        "Cooperative and competitive effective interactions jointly contribute to reproducing empirical static functional connectivity. Their strengths, signs and anatomical assignments form an interdependent fitted organization.",
    )
    add_heading(doc, "6.1 Connection to the landscape hypothesis", 2)
    add_body(
        doc,
        "The Phase 2 findings established a static-FC bridge to the broader proposal that cooperative and competitive interactions jointly shape a dynamical landscape. Phase 3 subsequently measured synchrony, metastability and continuous LEiDA landscape properties directly. It does not yet measure adaptive explore–exploit behaviour; that claim remains reserved for the explicit reversal-learning experiment."
    )

    add_heading(doc, "7. Limitations and safeguards", 1)
    for item in [
        "Single-subject scope: the results cannot yet be generalized across people.",
        "Simulation seeds are repeated stochastic runs, not independent biological observations.",
        "Frozen perturbations: signs and magnitudes were optimized jointly; shuffling one component creates a deliberately mismatched configuration and does not test all possible reoptimized alternatives.",
        "Effective-connectivity interpretation: negative generative weights represent model-inferred competitive influence and should not be equated with direct long-range inhibitory anatomy.",
        "Static outcome: FC correlation does not by itself characterize the temporal organization of brain states.",
        "Permutation resolution: with 100 shuffles, the smallest corrected one-sided permutation estimate is 1/101 = 0.0099.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Phase 3 handoff — Dynamical validation", 1)
    add_body(doc, "Status: core characterization complete.", bold_lead="Status:")
    add_body(
        doc,
        "Phase 3 will test whether the empirical, cooperative-only and signed models differ in how activity coordinates and changes through time. Metrics will be validated on synthetic signals before being interpreted on empirical or simulated BOLD data."
    )
    for item in [
        "Compute instantaneous phase for every regional signal.",
        "Compute the Kuramoto order parameter at every time point.",
        "Define synchrony as mean global phase coordination over time.",
        "Define metastability as temporal variability in global phase coordination.",
        "Validate the metrics on synchronized, independent, metastable and random synthetic signals.",
        "Compare empirical, cooperative-only and signed dynamics.",
        "Later add windowed FC, recurrent-state clustering, dwell time, occupancy and transition entropy.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Phase 3 gate",
        "The implemented metrics must correctly distinguish synthetic synchronized, independent, metastable and random signals before any biological interpretation is attempted.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "8.1 Measurement-instrument validation", 2)
    add_body(
        doc,
        "The first Phase 3 instrument was implemented in two layers: the Hilbert transform estimates instantaneous phase from narrow-band regional signals, and the Kuramoto order parameter summarizes global phase alignment at each timepoint. Synchrony is the temporal mean of R(t), and metastability is its population standard deviation."
    )
    add_result_table(
        doc,
        ["Synthetic regime", "Synchrony", "Metastability", "Expected behavior"],
        [
            ["Synchronized", "1.000", "≈ 0", "High, stable coordination"],
            ["Evenly dispersed", "≈ 0", "≈ 0", "Low, stable coordination"],
            ["Switching", "0.523", "0.376", "Alternates between aligned and dispersed states"],
            ["Random phase", "0.086", "0.044", "Low, irregular coordination"],
        ],
        [2600, 1700, 1800, 3260],
    )
    add_body(
        doc,
        "All seven predefined validation gates passed, including exact known-phase checks, recovery of a known sinusoidal phase offset and detection of high metastability in the switching regime. All 18 project tests passed. The measuring instrument therefore behaves correctly on these controlled synthetic cases."
    )
    add_figure(
        doc,
        "phase_dynamics_validation.png",
        "Figure 5. Phase-dynamics measurement validation on four synthetic coordination regimes.",
    )
    add_callout(
        doc,
        "Next methodological decision",
        "Before applying the instrument to BOLD data, define and verify the narrow-band filtering range and Hilbert-transform boundary trimming. Instantaneous phase is interpretable only for appropriately narrow-band signals.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "8.2 Empirical and frozen-model phase dynamics", 2)
    add_body(
        doc,
        "The paper and released code specify a second-order Butterworth human BOLD bandpass of 0.008–0.09 Hz. The identical detrending, zero-phase filtering, Hilbert phase extraction and Kuramoto analysis were applied to the empirical BOLD and to frozen cooperative-only and signed simulations across 30 matched noise seeds."
    )
    add_result_table(
        doc,
        ["Signal", "Mean KOP", "Maximum KOP", "Metastability: SD(KOP)"],
        [
            ["Empirical BOLD", "0.279", "0.704", "0.138"],
            ["Cooperative-only mean", "0.395", "0.775", "0.170"],
            ["Signed mean", "0.302", "0.655", "0.141"],
        ],
        [3000, 1900, 2100, 2360],
    )
    add_body(
        doc,
        "The signed model was closer to empirical mean synchrony in 30/30 matched seeds and closer in metastability in 25/30 seeds. The cooperative-only model was excessively coordinated on average and showed excessive variability in coordination. Competitive interactions moderated both quantities toward the empirical values."
    )
    add_body(
        doc,
        "Maximum synchrony was less decisive: the signed model was closer for 16/30 seeds and tended to undershoot the empirical maximum, whereas the cooperative-only model tended to overshoot it. Thus, the signed model more clearly improves mean coordination and metastability than the most extreme synchronization event."
    )
    add_figure(
        doc,
        "empirical_phase_dynamics.png",
        "Figure 6. Empirical KOP time series and frozen-model metastability and maximum-synchrony distributions across 30 matched seeds.",
    )
    add_callout(
        doc,
        "Boundary sensitivity",
        "The qualitative conclusion was unchanged when excluding 0, 20, 50 or 100 samples from each edge. The primary analysis uses no trimming because the paper does not report it; the alternate trims serve as sensitivity checks.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "8.3 From discrete states to a continuous landscape", 2)
    add_body(
        doc,
        "Windowed-FC k-means was first validated on synthetic signals and stress-tested across 768 combinations of noise, state duration and window size. Although it recovered clear synthetic states within known operating bounds, empirical solutions did not identify one uniquely stable number of states. Projective LEiDA corrected the arbitrary eigenvector-sign ambiguity, but no empirical two- to ten-state solution passed every strict stability gate. Discrete states were therefore retained as descriptive summaries rather than treated as ground truth."
    )
    add_body(
        doc,
        "A continuous LEiDA representation became the primary instrument. It quantifies repertoire dispersion, effective dimension, distance from the central axis, movement speed, speed variability and nonlocal recurrence without forcing every time point into a hard state. Synthetic fixed, switching, circular and wandering trajectories passed all six predefined validation gates."
    )
    add_body(
        doc,
        "Across 30 frozen simulations, the signed model was more often closer to empirical dispersion (27/30), effective dimension (28/30) and central distance (30/30). The cooperative-only model was more often closer in mean speed (22/30) and speed variability (19/30). This motivated a provisional distinction between the accessible landscape and movement through it."
    )

    add_heading(doc, "8.4 Cooperative–competitive gain landscape", 2)
    add_body(
        doc,
        "Cooperative and competitive weights were separated while preserving their fitted anatomical placement, then multiplied by independent gains. Competition alone collapsed the repertoire, while removing competition from the signed solution yielded incomplete static and dynamical reproduction. Balanced gain was non-monotonic: low gain produced excessive effective dimension, the original gain approached empirical geometry and high gain compressed the landscape."
    )
    add_callout(
        doc,
        "Static and dynamic objectives diverge",
        "The highest tested balanced gain produced the strongest static FC fit but poorer landscape geometry. A single FC objective therefore cannot establish overall dynamical realism.",
        fill=LIGHT_GRAY,
    )

    doc.add_page_break()
    add_heading(doc, "8.5 Noise and bifurcation: independent confirmation", 2)
    add_body(
        doc,
        "The signed connectivity and sign gains were frozen. A coarse sweep showed little sensitivity to noise over the tested range but strong, coupled effects of the Hopf bifurcation parameter. An exploratory refinement selected a = -0.025 with noise = 0.004. This candidate and four controls were then evaluated on 30 new paired seeds (102–131), with no reuse of the selection seeds."
    )
    add_result_table(
        doc,
        ["Measure", "Baseline", "Candidate", "Candidate closer"],
        [
            ["FC correlation", "0.670", "0.656", "0/30"],
            ["Mean speed", "0.208", "0.220", "26/30"],
            ["Central distance", "1.156", "1.185", "25/30"],
            ["Effective dimension", "13.39", "15.49", "10/30"],
            ["Recurrence distance", "0.943", "0.980", "1/30"],
        ],
        [3000, 1900, 1900, 2560],
    )
    add_body(
        doc,
        "The candidate robustly improved mean speed and central distance, but worsened effective dimension, recurrence and static FC. The bifurcation-only control at a = -0.025 and noise = 0.001 was nearly identical to the candidate, demonstrating that the effect was driven primarily by the bifurcation change rather than increased noise."
    )
    add_callout(
        doc,
        "Hypothesis revision",
        "Connectivity and cooperative–competitive balance strongly constrain the landscape, while the bifurcation parameter strongly affects movement. However, the bifurcation parameter also reshapes the landscape: geometry and kinetics are coupled rather than independently adjustable dials.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "8.6 Phase 3 decision", 2)
    add_body(
        doc,
        "The negative result is informative and independently reproducible: there was no clean kinetic rescue that preserved every geometrical and static property. Phase 3 is sufficiently mature to proceed. The next stage will define and validate an explicit reversal-learning task before comparing connectome variants, preventing task design from being tuned to favor the signed model."
    )

    add_heading(doc, "8.7 Frozen EEG reversal-learning confirmation", 2)
    add_body(
        doc,
        "A switch-locked EEG analysis was frozen before inspecting confirmation participants sub-s2, sub-s3 and sub-s4 from OpenNeuro ds004295. The primary endpoint asked whether theta phase-pattern reconfiguration on negative-feedback reward trials was greater during the first ten trials after reversal than during the ten stable trials before reversal. A participant required at least five clean trials in both periods, and the directional gate required a positive effect in at least two participants."
    )
    add_result_table(
        doc,
        ["Participant", "Stable trials", "Early trials", "Early − stable", "Decision"],
        [
            ["sub-s2", "11", "10", "−0.002875", "Negative direction"],
            ["sub-s3", "4", "13", "Not evaluated", "Below frozen minimum"],
            ["sub-s4", "12", "11", "+0.006147", "Positive direction"],
        ],
        [1500, 1600, 1500, 2100, 2660],
    )
    add_callout(
        doc,
        "Frozen gate result — did not pass",
        "Only one of the two evaluable participants showed the predicted positive direction. The threshold was not lowered after seeing the data. This is a failed technical and directional confirmation, not evidence that reversal learning has no neural reconfiguration.",
        fill=LIGHT_GRAY,
    )
    add_body(
        doc,
        "The supporting checks were largely successful. All six EEG files passed integrity verification, artifact-rejection imbalance remained below 20 percentage points for all participants, and the frontal-midline reward-theta positive control passed in two of three participants. Participant sub-s3 was excluded mechanically because only four clean stable-pre negative-reward trials remained."
    )
    add_body(
        doc,
        "The primary values clustered near 0.96–0.97, revealing a possible ceiling or scaling weakness in the instantaneous phase-pattern metric. The next work is therefore an explicitly exploratory instrument audit using synthetic known-state signals and surrogate EEG. Alternative estimators must be validated and frozen before testing additional untouched participants; the failed gate will remain unchanged as the record of this attempt."
    )
    add_figure(
        doc,
        "reversal_eeg_confirmation.png",
        "Figure 7. Frozen EEG confirmation result. Participant sub-s3 was not plotted because it did not meet the prespecified stable-period trial minimum.",
    )

    add_heading(doc, "8.8 Post-failure estimator comparison", 2)
    add_body(
        doc,
        "The near-ceiling result prompted a formally separated exploratory instrument study. Four candidate estimators were tested on 64-channel synthetic theta signals with five known reconfiguration levels: equal-window phase-locking-value matrices, weighted phase-lag-index matrices, projective LEiDA distributions and Riemannian covariance geometry. Equal 1.3-second pre/post windows were used throughout."
    )
    add_body(
        doc,
        "The central mitigation was to ask how much more the observed configuration changed than expected from ordinary oscillator evolution. Every synthetic observation was therefore paired with a zero-change counterfactual preserving the same regional frequencies, phase noise, sensor noise and volume mixing. The score was observed distance minus matched expected-drift distance."
    )
    add_result_table(
        doc,
        ["Estimator", "Strength", "Failure or boundary", "Decision"],
        [
            ["Equal-window PLV", "Drift-stable", "Blind to fixed phase-topology change", "Reject for primary question"],
            ["Equal-window wPLI", "Lagged coupling; mixing-resistant", "Lost sensitivity with frequency heterogeneity", "Retain as secondary"],
            ["Projective LEiDA", "Sensitive under noise and mixing", "Short-window drift dominated distributions", "Requires explicit drift model"],
            ["Riemannian covariance", "Strongest overall ordering", "Broader than phase; small realistic residual", "Retain as leading parallel candidate"],
        ],
        [1900, 2300, 2700, 2460],
    )
    add_body(
        doc,
        "Riemannian covariance was the strongest overall candidate. It remained perfectly ordered under isolated frequency drift and achieved a Spearman ordering of 0.90 under combined realistic confounds, although its realistic dynamic range was small and its final two levels were not strictly monotonic. No tested raw distance was sufficiently robust to become the new primary endpoint."
    )
    add_callout(
        doc,
        "Methodological result",
        "Equal-duration windows alone do not solve the construct problem, and subtracting an average null is insufficient. Reversal-related change must be evaluated against a trial-specific prediction of ordinary oscillator evolution.",
        fill=LIGHT_GRAY,
    )
    add_figure(
        doc,
        "eeg_reconfiguration_estimator_comparison.png",
        "Figure 8. Drift-calibrated comparison of four EEG reconfiguration estimators across known synthetic changes and increasingly realistic confounds.",
    )
    add_body(
        doc,
        "The next instrument will estimate each channel’s pre-feedback phase trajectory, extrapolate its expected post-feedback evolution and quantify the residual between observed and predicted organization. Stable pseudo-events will establish prediction error in the absence of reversal. Residual phase topology, residual LEiDA, Riemannian covariance and wPLI will then be compared before a revised estimator is frozen for untouched participants."
    )

    add_heading(doc, "9. Update protocol", 1)
    add_body(
        doc,
        "After each substantive milestone, update this report with: the exact question, the controlled manipulation, preserved quantities, simulation and biological sample sizes, primary numerical results, a plain-language interpretation, limitations, output paths and the next decision. Clearly separate confirmatory tests from exploratory analyses."
    )

    add_heading(doc, "10. Reproducibility index", 1)
    add_result_table(
        doc,
        ["Artifact", "Location"],
        [
            ["Project roadmap", "PROJECT_OUTLINE.md"],
            ["Upstream version record", "configs/upstream_version.txt"],
            ["Upstream modification notes", "configs/upstream_patch_notes.md"],
            ["Core optimization results", "results/single_subject_optimization/"],
            ["Frozen stochastic evaluation", "results/frozen_gc_evaluation/"],
            ["Whole-weight placement null", "results/weight_placement_null/"],
            ["Sign-specific organization", "results/sign_specific_organization/"],
            ["Anatomical sign-map null", "results/sign_map_null/"],
            ["Continuous LEiDA landscape", "results/leida_landscape_evaluation/"],
            ["Gain confirmation", "results/cooperative_competitive_gain_confirmation/"],
            ["Noise/bifurcation confirmation", "results/noise_bifurcation_confirmation/"],
            ["Frozen EEG protocol", "docs/EEG_SWITCHING_ANALYSIS_PROTOCOL.md"],
            ["Frozen EEG confirmation result", "docs/EEG_CONFIRMATION_RESULTS.md"],
            ["EEG confirmation outputs", "results/reversal_eeg_confirmation/"],
            ["EEG estimator comparison", "docs/EEG_ESTIMATOR_COMPARISON.md"],
            ["Estimator comparison outputs", "results/eeg_reconfiguration_estimator_comparison/"],
            ["Analysis code and automated tests", "scripts/ and tests/"],
        ],
        [3600, 5760],
    )

    # The table helper adds a spacer paragraph. Remove the final spacer so a
    # full-page reproducibility table does not create a blank trailing page.
    if doc.paragraphs and not doc.paragraphs[-1].text:
        paragraph = doc.paragraphs[-1]._element
        paragraph.getparent().remove(paragraph)

    properties = doc.core_properties
    properties.title = "Competitive Connectomes and Adaptive Flexibility — Research Progress Report"
    properties.subject = "Living record of computational methods, results and interpretation"
    properties.author = "Shriya Sai"
    properties.keywords = "Hopf model, effective connectivity, cooperation, competition, NeuroAI"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
