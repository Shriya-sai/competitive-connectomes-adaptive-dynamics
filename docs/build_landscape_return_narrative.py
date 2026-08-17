"""Build the narrative tracing the project back to the explore-exploit paper."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "From_Competitive_Connectomes_Back_to_Explore_Exploit.docx"
FIGURES = ROOT / "figures"

BLUE = "2E5D7B"
DARK = "173A52"
PALE = "E8EEF5"
GRAY = "68717A"
GOLD = "A77A2D"
INK = "1A1A1A"


def font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def margins(cell, top=100, bottom=100, start=140, end=140):
    properties = cell._tc.get_or_add_tcPr()
    element = properties.first_child_found_in("w:tcMar")
    if element is None:
        element = OxmlElement("w:tcMar")
        properties.append(element)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = element.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            element.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, value in zip(row.cells, widths, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            cell.width = Inches(value / 1440)
            node = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")


def configure(doc):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.194)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.208


def body(doc, text, lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    if lead and text.startswith(lead):
        font(paragraph.add_run(lead), bold=True)
        font(paragraph.add_run(text[len(lead):]))
    else:
        font(paragraph.add_run(text))
    return paragraph


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    font(paragraph.add_run(text))


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    font(paragraph.add_run(f"{label}: "), bold=True, color=DARK)
    font(paragraph.add_run(text))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def comparison_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    geometry(table, [2300, 3530, 3530])
    for cell, label in zip(table.rows[0].cells, ("Concept", "Original framework", "Computational result"), strict=True):
        shade(cell, PALE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(paragraph.add_run(label), size=9.5, bold=True, color=DARK)
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values, strict=True)):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font(paragraph.add_run(value), size=9.5, bold=index == 0)
    geometry(table, [2300, 3530, 3530])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figure(doc, filename, caption, width=6.25):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(3)
    caption_paragraph.paragraph_format.space_after = Pt(9)
    font(caption_paragraph.add_run(caption), size=9, italic=True, color=GRAY)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    font(run, size=9, color=GRAY)


def build():
    doc = Document()
    configure(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    font(header.add_run("COMPETITIVE CONNECTOMES AND ADAPTIVE FLEXIBILITY"), size=8.5, bold=True, color=GRAY)
    page_number(section.footer.paragraphs[0])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(86)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    font(kicker.add_run("RESEARCH NARRATIVE"), size=10, bold=True, color=GOLD)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(9)
    font(title.add_run("How the Computational Project\nLed Back to the Explore–Exploit Framework"), size=27, bold=True, color=DARK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    font(subtitle.add_run("From reproducing competitive Hopf dynamics to measuring the geometry and kinetics of an adaptive brain landscape"), size=14, color=BLUE)
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(metadata.add_run(f"Competitive Connectomes and Adaptive Flexibility  |  {date(2026, 8, 6).strftime('%d %B %Y')}"), size=10, color=GRAY)
    doc.add_page_break()

    heading(doc, "Executive synthesis")
    body(doc, "This project began as a methodological reconstruction of Luppi et al.’s cooperative–competitive Hopf model. It did not begin by fitting a new explore–exploit theory to data. The path back to that theory emerged incrementally: first by establishing that signed connectivity improves functional and dynamical realism; then by showing that anatomical placement of both cooperative and competitive interactions matters; and finally by replacing unstable discrete state labels with a continuous geometric description of the system’s phase-locking repertoire.")
    callout(doc, "Central result", "The signed model reproduced the empirical breadth, dimensionality, and central geometry of the accessible phase-locking landscape much better than the cooperative-only model. The cooperative-only model, however, reproduced moment-to-moment trajectory speed more closely. This separates the geometry of what the system can access from the kinetics of how it moves through that space.")
    body(doc, "That separation returns directly to the paper’s organizing intuition: interactions can shape an adaptive landscape, while other dynamical controls regulate exploration, persistence, escape, and transition within it. The result is not proof of the full cognitive theory, but it is a concrete computational bridge from its conceptual language to measurable whole-brain dynamics.")

    heading(doc, "1. The conceptual starting point")
    body(doc, "The original paper treated explore–exploit behavior as a dynamical problem rather than a simple choice between novelty and repetition. Its deeper claim was that adaptive cognition depends on a structured landscape of possible configurations and on the system’s ability to move through that landscape. Cooperation and competition were proposed as complementary forces: neither was merely beneficial or harmful, and flexibility arose from their coordinated balance.")
    bullet(doc, "Landscape geometry concerns which configurations are available, separated, connected, or favored.")
    bullet(doc, "Exploration concerns movement across configurations and access to alternatives.")
    bullet(doc, "Exploitation concerns persistence near a useful configuration or basin.")
    bullet(doc, "Adaptive flexibility concerns changing persistence or movement when conditions change.")
    body(doc, "At this stage these were conceptual commitments. The computational project was designed first to understand Luppi et al.’s model faithfully, not to manufacture confirmation of them.")

    heading(doc, "2. Reconstructing the Luppi methodology")
    body(doc, "The reconstruction used the released single-subject public dataset: 100 regional BOLD signals, a 100 × 100 structural connectivity matrix, and TR = 0.72 seconds. Regional oscillation frequencies were extracted in Python and matched the released MATLAB reference exactly. The C++ Hopf simulator was then built and exercised with frozen fitted cooperative-only and signed generative connectivity matrices.")
    body(doc, "Across 30 matched stochastic seeds, the signed model reproduced empirical functional connectivity more accurately in every run. Mean FC correlation increased from 0.482 for cooperative-only simulations to 0.677 for signed simulations. This established a robust improvement in static fit, but it did not yet demonstrate the landscape interpretation.")
    figure(doc, "frozen_gc_stochastic_evaluation.png", "Figure 1. Frozen-model functional-connectivity evaluation across matched stochastic seeds.")

    heading(doc, "3. Why both signs and their locations mattered")
    body(doc, "A sequence of frozen perturbation tests then asked whether performance came merely from adding negative numbers or from the organized placement of both signs. Whole-weight placement shuffling collapsed FC similarity. Shuffling negative and magnitude-matched positive connections separately degraded fit in both cases. Finally, permuting reciprocal sign patterns while holding every anatomical connection magnitude fixed also destroyed the fitted advantage.")
    callout(doc, "Interpretation", "The result was not that competition alone carries the system’s organization. Cooperative and competitive placements jointly encode location-specific information, and the fitted sign map is interdependent with structural topology and connection magnitudes.")
    figure(doc, "sign_specific_organization.png", "Figure 2. Sign-specific organization tests show that disrupting either cooperative or competitive placement reduces empirical fit.")

    heading(doc, "4. From static fit to global dynamics")
    body(doc, "Kuramoto order parameter analysis compressed the 100 regional phases into one instantaneous measure of global synchrony. Its mean described average coordination, and its temporal standard deviation described metastability. The signed model was closer to empirical mean synchrony in all 30 simulations and closer to empirical metastability in 25 of 30. Cooperation alone produced excessive average coordination and excessive fluctuation in coordination.")
    body(doc, "This supported a coordination–release interpretation, but KOP could not reveal which regional configurations produced the same global value. A new method was required to examine the landscape’s internal structure.")
    figure(doc, "empirical_phase_dynamics.png", "Figure 3. Empirical and frozen-model phase dynamics: signed interactions moderate excessive cooperative coordination toward empirical values.")

    heading(doc, "5. The discrete-state detour—and why failure mattered")
    body(doc, "Sliding-window functional connectivity and k-means were validated successfully on synthetic signals with known switches. Under empirical BOLD, however, ordinary clustering was unstable across initializations and window sizes. Consensus clustering recovered reproducible organization at longer windows, but six, seven, and eight states all passed, leaving the correct granularity ambiguous.")
    body(doc, "LEiDA was introduced to remove the arbitrary window. A first implementation failed synthetic validation because eigenvectors v and −v represent the same axis. That failure exposed an inappropriate Euclidean assumption. A projective clustering method was implemented, after which synthetic state recovery was essentially perfect under substantial phase noise.")
    body(doc, "Empirical LEiDA still did not yield a uniquely robust multistate atlas. A two-regime candidate was strong but fell slightly below the prespecified mean block-stability threshold. This did not imply that the data were random. It indicated that forcing hard state boundaries was not the most defensible primary representation for this single recording.")
    callout(doc, "Methodological decision", "The project stopped asking for a single correct number of discrete states and moved to a continuous, sign-invariant LEiDA landscape. The failed state searches therefore redirected the analysis rather than being discarded.", fill="F4F6F9")

    heading(doc, "6. Measuring the continuous landscape")
    body(doc, "Each timepoint was represented by the dominant regional phase-locking axis. Projective angular geometry treated v and −v as identical. Five complementary properties were then defined without clustering: pairwise repertoire dispersion, effective dimension, trajectory speed, distance from the central axis, and nearest recurrence outside a short temporal neighborhood. A sixth measure captured variability in speed.")
    body(doc, "The instrument was validated on four known synthetic regimes. A fixed trajectory had zero spread and one effective dimension; a switching trajectory occupied two axes; a circular trajectory repeatedly returned through a smooth path; and a random wandering trajectory had high dispersion, high speed, weak recurrence, and approximately twenty effective dimensions. All validation gates passed.")
    figure(doc, "leida_landscape_validation.png", "Figure 4. Continuous LEiDA metrics correctly distinguish fixed, switching, recurrent circular, and wandering trajectories.")

    heading(doc, "7. The result that returned to the paper")
    body(doc, "The frozen empirical comparison used 30 matched seeds and no state fitting. The signed model was closer to empirical repertoire dispersion in 27 of 30 runs, effective dimension in 28 of 30, central distance in all 30, and nearest recurrence distance in 20 of 30. The strongest geometric separation occurred in effective dimensionality and central distance.")
    comparison_table(doc, [
        ("Repertoire breadth", "A flexible system should access a sufficiently broad set of configurations.", "Signed mean dispersion 1.341 versus empirical 1.351; cooperative-only 1.292."),
        ("Dimensional richness", "The landscape should support multiple independent directions, not collapse to one dominant route.", "Signed effective dimension 12.71 versus empirical 13.76; cooperative-only 9.70."),
        ("Departure from center", "Exploration requires leaving a dominant or habitual configuration.", "Signed central distance 1.147 versus empirical 1.205; cooperative-only 1.036. Signed was closer in 30/30 runs."),
        ("Movement kinetics", "Flexibility also depends on transition speed and variability, not only available destinations.", "Cooperative-only mean speed 0.227 versus empirical 0.225; signed 0.206. Cooperative-only was closer in 22/30 runs."),
    ])
    figure(doc, "leida_landscape_evaluation.png", "Figure 5. Continuous empirical and frozen-model LEiDA landscapes. Dashed lines mark empirical values.")

    heading(doc, "8. The bridge back to explore–exploit")
    body(doc, "The project therefore arrived at a separation that closely mirrors the paper’s architecture. Cooperative and competitive connectivity strongly influenced the geometry of the accessible repertoire—its breadth, dimensionality, and distance from a dominant center. Yet the same signed model did not optimally reproduce movement speed or its variability.")
    callout(doc, "Emergent hypothesis", "Cooperative and competitive connectivity may shape the topology of the accessible dynamical landscape, whereas noise strength, bifurcation proximity, delays, and separately scaled interaction gains may regulate exploration rate, persistence, escape, and transition kinetics within that landscape.")
    body(doc, "This statement must be read carefully. The first clause is supported here at the level of a single-subject fitted Hopf model: changing the signed interaction architecture altered continuous landscape geometry. The second clause is a mechanistic hypothesis motivated by the remaining speed mismatch. Noise, bifurcation, delays, and separate gains have not yet been systematically varied.")
    body(doc, "The intellectual return is therefore genuine but bounded. The conceptual paper anticipated that landscape structure and movement through it are related but distinguishable. The computational project independently encountered the same distinction after method reconstruction, falsifiable perturbation tests, failed discrete-state searches, and continuous geometric measurement.")

    heading(doc, "9. What has—and has not—been established")
    heading(doc, "Established in this project", 2)
    bullet(doc, "The signed frozen model reproduces static empirical FC better than the cooperative-only model across 30 matched seeds.")
    bullet(doc, "The anatomical placement of both cooperative and competitive interactions contains location-specific information.")
    bullet(doc, "Signed dynamics reproduce empirical mean synchrony and metastability more closely.")
    bullet(doc, "The signed model reproduces continuous landscape breadth, dimensionality, and central geometry more closely.")
    bullet(doc, "The cooperative-only model better reproduces instantaneous trajectory speed and speed variability in the present parameterization.")
    heading(doc, "Not yet established", 2)
    bullet(doc, "That the same effects generalize across biological subjects, species, tasks, or cognitive states.")
    bullet(doc, "That the measured phase-locking geometry directly implements psychological exploration or exploitation.")
    bullet(doc, "That life experience creates the fitted signed structural or effective connections.")
    bullet(doc, "That a unique set of discrete brain states exists in this recording.")
    bullet(doc, "Which parameters causally control transition speed, persistence, or escape.")

    heading(doc, "10. The next falsifiable experiments")
    body(doc, "The return to the paper should now guide experiments rather than serve as a retrospective metaphor. The next phase should vary putative geometry and kinetics controls separately while preserving held-out evaluation.")
    bullet(doc, "Sweep cooperative and competitive gains independently and map changes in repertoire dispersion, dimensionality, central distance, speed, and recurrence.")
    bullet(doc, "Vary noise and bifurcation parameters while holding fitted connectivity fixed to test whether kinetics can be corrected without destroying geometry.")
    bullet(doc, "Introduce delays only after the simpler parameter sweeps establish identifiable effects.")
    continued = heading(doc, "Next experiments (continued)", 2)
    continued.paragraph_format.page_break_before = True
    bullet(doc, "Repeat the analysis on additional public subjects and learn empirical targets from training subjects before testing held-out individuals.")
    bullet(doc, "Connect landscape measures to reversal-learning or explore–exploit behavior only when an appropriate public task dataset is available.")
    callout(doc, "Decision rule", "The framework gains support only if geometry-related manipulations and kinetics-related manipulations show partially separable, reproducible effects. If all parameters affect all measurements indiscriminately, the proposed decomposition must be revised.", fill="F4F6F9")

    heading(doc, "Sources and project artifacts")
    body(doc, "Primary paper: Luppi, A. I. et al. Competitive interactions shape mammalian brain network dynamics and computation. Nature Neuroscience 29, 915–933 (2026).")
    body(doc, "Project evidence: frozen GC evaluation; sign-specific and sign-map null analyses; empirical phase-dynamics analysis; recurrent-state validation and stress tests; empirical consensus analyses; LEiDA state validation; empirical LEiDA selection; and continuous LEiDA landscape validation and evaluation. All numerical values reported here are drawn from the saved JSON and CSV outputs in the project results directory.")

    doc.core_properties.title = "How the Computational Project Led Back to the Explore–Exploit Framework"
    doc.core_properties.subject = "Competitive connectomes, adaptive flexibility, and continuous brain landscapes"
    doc.core_properties.author = "Shriya Sai"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
